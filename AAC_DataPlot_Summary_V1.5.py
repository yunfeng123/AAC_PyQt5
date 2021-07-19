import pandas as pd
from tkinter import *
from tkinter import filedialog
import tkinter.messagebox
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import main_window_action as mw
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow


# def run_report():
#     pic_path = v.get()
#     list_pic = os.listdir(pic_path)
#
#     # pic_name 为图片名称拆分组成的DataFrame（分组名,Config,图类），columns为图片名称
#     pic_name = pd.DataFrame()
#     for i in list_pic:
#         if os.path.isfile(pic_path + '\\' + i):
#             list_temp = i.split('-')
#             list_temp[2] = int(list_temp[2])
#             series_temp = pd.Series(list_temp[0:4])
#             pic_name = pd.concat([pic_name, series_temp], axis=1)
#         else:
#             list_pic.remove(i)
#     pic_name.columns = list_pic
#
#     pic_name.sort_values(by=[3, 2, 0, 1], axis=1, inplace=True)  # 排序，先按照station，再按照图类别，然后是组的类别，最后是Config
#     # pic_name.to_csv('pic_name.csv')
#
#     # sections对应文档中的“节”
#     document = Document()
#     sec = document.sections[0]
#
#     # 页眉设置
#     sec.header_distance = Cm(0.3)  # 页眉距离顶端距离
#     paragraph = sec.header.paragraphs[0]
#     run_header = paragraph.add_run('TestDataPlotSummary developed by AAC Test Team')
#     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     run_header.font.name = u'Calibri'
#     run_header.font.size = Pt(8)
#     run_header._element.rPr.rFonts.set(qn('w:eastAsia'), u'Calibri')
#     run_header.font.color.rgb = RGBColor(0, 0, 0)
#     run_header.font.italic = True
#
#     # 以下依次设置左、右、上、下页面边距
#     distance1 = Inches(0.2)
#     distance2 = Inches(0)
#     sec.left_margin = distance1
#     sec.right_margin = distance1
#     sec.top_margin = distance2
#     sec.bottom_margin = distance2
#
#     # 设置页面的宽度和高度(A4)
#     sec.page_width = Inches(11.69)
#     sec.page_height = Inches(8.27)
#
#     flag_index = pd.Series(['', '', '', ''])  # 上一条运行的图片分类，用于下一次检测是否标题，换行等
#     config_qty_tem = 1  # 用于检测是否大于3个config了，换页添加标题
#     loop_count = 0
#
#     #   循环逻辑：（1）Type不一样，Type标题+Config标题+图片，新建Par-RUN；（2）Config不一样，Config标题+图片，新建Rar-RUN；（3）正常绘图
#     for i in pic_name.columns:
#
#         # 主标题设置
#         if pic_name[i][3] != flag_index[3]:
#             if loop_count > 0:  # 第一页不换页
#                 document.add_page_break()
#             page_title = document.add_heading("", level=1)
#             page_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             run_pt = page_title.add_run('\n\n\n' + 'TestData Plot Summary Report' + '\n' + pic_name[i][3])
#             run_pt.font.name = u'Helvetica'
#             run_pt.font.size = Pt(40.5)
#             run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
#             run_pt.font.color.rgb = RGBColor(0, 0, 0)
#             run_pt.font.bold = True
#
#         if pic_name[i][0] != flag_index[0] or (pic_name[i][1] != flag_index[1] and config_qty_tem % 3 == 0):  # 图类别
#             document.add_page_break()
#             page_title = document.add_heading('', level=1)
#             page_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#             run_pt = page_title.add_run(pic_name[i][0])
#             run_pt.font.name = u'Helvetica'
#             run_pt.font.size = Pt(17)
#             run_pt._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
#             run_pt.font.color.rgb = RGBColor(0, 0, 0)
#             run_pt.italic = True
#             run_pt.font.bold = True
#
#             page_title_2nd = document.add_heading('', level=2)
#             page_title_2nd.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#             run_ct = page_title_2nd.add_run(pic_name[i][1])
#             run_ct.font.name = u'Helvetica'
#             run_ct.font.size = Pt(12)
#             run_ct._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
#             run_ct.font.color.rgb = RGBColor(0, 0, 0)
#
#             pr_pic = document.add_paragraph()
#             run_pict = pr_pic.add_run()
#
#             run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
#             flag_index = pic_name[i]
#
#             config_qty_tem = 1
#
#
#         elif pic_name[i][1] != flag_index[1]:  # Config
#
#             page_title_2nd = document.add_heading('', level=2)
#             page_title_2nd.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#             run_ct = page_title_2nd.add_run(pic_name[i][1])
#             run_ct.font.name = u'Helvetica'
#             run_ct.font.size = Pt(12)
#             run_ct._element.rPr.rFonts.set(qn('w:eastAsia'), u'Helvetica')
#             run_ct.font.color.rgb = RGBColor(0, 0, 0)
#
#             pr_pic = document.add_paragraph()
#             run_pict = pr_pic.add_run()
#
#             run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
#             flag_index = pic_name[i]
#
#             config_qty_tem = config_qty_tem + 1
#
#         else:
#             run_pict.add_picture(pic_path + '\\' + i, height=Inches(2))
#             flag_index = pic_name[i]
#
#         loop_count = loop_count + 1
#
#     config_name = entry_config.get()
#     report_path = os.path.split(pic_path)[0] + '/' + config_name + f'_Data Summary Report_WD_{now_date}.docx'
#     document.save(report_path)
#     text_info_line = text_info.index(END).split('.')[0]
#     text_info.insert(END, 'Finished Report in ' + report_path + '\n')
#     text_info.tag_add('tag2', str(int(text_info_line) - 1) + '.0', text_info_line + '.0')  # Finish 打印突出显示
#     text_info.tag_config('tag2', background='Cyan1', font=('Times', 10))
#     text_info.see(END)
#     text_info.update()


# if now_date > overdue_date:
#     tkinter.messagebox.askokcancel(title='Error', message='License Expired !')
#     exit()

if __name__ == '__main__':
    # 创建QApplication类的实例
    app = QApplication(sys.argv)
    # 创建一个主窗口
    mainWindow = QMainWindow()
    # 创建Ui_MainWindow的实例
    ui =mw.mainwindow_action()
    # 调用setupUi在指定窗口(主窗口)中添加控件
    ui.setupUi(mainWindow)
    # 显示窗口
    mainWindow.show()
    # 进入程序的主循环，并通过exit函数确保主循环安全结束
    sys.exit(app.exec_())
